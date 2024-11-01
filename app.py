import requests
from bs4 import BeautifulSoup
from docx import Document
from io import BytesIO
import streamlit as st
from urllib.parse import urlparse
import re

# Optimized function to scrape and save content to DOCX
def scrape_and_save_to_docx(url):
    headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:102.0) Gecko/20100101 Firefox/102.0"}
    try:
        response = requests.get(url, headers=headers, timeout=5)
        response.raise_for_status()  # Raise an exception for unsuccessful requests
    except requests.exceptions.RequestException as e:
        if response.status_code == 403:
            st.error(f"Error: Access denied for URL (403 Forbidden).")
        elif response.status_code == 429:
            st.error(f"Error: Too many requests sent too quickly (429 Too Many Requests).")
        else:
            st.error(f"Error fetching webpage: {e}")
        return None

    # Parse the webpage content
    soup = BeautifulSoup(response.content, 'lxml')  # Switched to 'html.parser' for speed (over lxml)
    doc = Document()

    # Extract and clean domain name for DOCX heading and file naming
    domain = urlparse(url).netloc or url
    valid_filename = re.sub(r'[^a-zA-Z0-9_-]', '_', domain)

    doc.add_heading(f'Content from {domain}', level=1)

    # Direct tag handlers mapped to respective functions for speed
    def handle_tag(tag):
        tag_map = {
            'h1': lambda t: doc.add_heading(t.get_text(strip=True), level=1),
            'h2': lambda t: doc.add_heading(t.get_text(strip=True), level=2),
            'h3': lambda t: doc.add_heading(t.get_text(strip=True), level=3),
            'h4': lambda t: doc.add_heading(t.get_text(strip=True), level=4),
            'p': lambda t: doc.add_paragraph(t.get_text(strip=True)),
            'ul': lambda t: [doc.add_paragraph(f"• {li.get_text(strip=True)}", style='List Bullet') for li in t.find_all('li')],
            'ol': lambda t: [doc.add_paragraph(f"{i+1}. {li.get_text(strip=True)}", style='List Number') for i, li in enumerate(t.find_all('li'))],
            'img': lambda t: doc.add_paragraph(f"[Image Placeholder: {t.get('src', 'Image URL')}]"),
            'audio': lambda t: doc.add_paragraph(f"[Audio Placeholder: {t.get('src', 'Audio URL')}]"),
            'video': lambda t: doc.add_paragraph(f"[Video Placeholder: {t.get('src', 'Video URL')}]")
        }
        if tag.name in tag_map:
            tag_map[tag.name](tag)

    # Efficiently search and handle tags within the <body> for faster processing
    for tag in soup.body.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'ul', 'ol', 'img', 'audio', 'video']):
        handle_tag(tag)

    # Save to DOCX in memory (BytesIO)
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    return doc_io, f"{valid_filename}.docx"

# Function to display DOCX content on the screen
def display_docx_content(docx_io):
    document = Document(docx_io)
    return "\n".join([para.text for para in document.paragraphs])

# Streamlit app
st.title("Rapid Web Scraper to DOCX")

# Input for URL
url = st.text_input("Enter the URL of the webpage to scrape")

# Button to scrape and generate DOCX file
if st.button("Generate and View DOCX"):
    if url:
        with st.spinner("Scraping content and generating DOCX..."):
            docx_io, file_name = scrape_and_save_to_docx(url)
            
            if docx_io:
                st.success(f"DOCX file '{file_name}' generated successfully!")

                # Display the DOCX content on the screen
                docx_content = display_docx_content(docx_io)
                st.text_area("DOCX Content", docx_content, height=300)

                # Download button for the DOCX file
                st.download_button(
                    label="Download DOCX",
                    data=docx_io,
                    file_name=file_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    else:
        st.error("Please enter a valid URL.")
